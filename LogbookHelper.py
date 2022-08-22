import docx
import datetime
import re
from docx.shared import Pt
import os

os.chdir('D:')
os.chdir("!Research Asst")

document = docx.Document('August Logbook 130.docx')

#Changing dates on initial page to current month
document.tables[0].cell(8,1).text=datetime.datetime.now().strftime('%B')+' '+re.split(' ',document.tables[0].cell(8,1).text)[1]
document.tables[0].cell(8,1).paragraphs[0].paragraph_format.space_after = Pt(1)
list_sign=re.split('Date: ',document.tables[1].cell(0,0).text)
list_sign[-1]=datetime.datetime.now().strftime('%d %B %G')+'\n'
document.tables[1].cell(0,0).text='Date: '.join(list_sign)

list_sign=re.split('Date: ',document.tables[1].cell(0,1).text)
list_sign[-1]=datetime.datetime.now().strftime('%d %B %G')+'\n'
document.tables[1].cell(0,1).text='Date: '.join(list_sign)

if len(document.tables)==2:
    day_num=1
else:
    day_num=int(re.split('Day Number: ',document.tables[-1].cell(0,0).text)[1])+1
document.add_paragraph()  
document.add_table(7,1)
    
document.tables[-1].cell(0,0).paragraphs[0].add_run('Day Number: %i'%(day_num)).bold=True
document.tables[-1].cell(1,0).paragraphs[0].add_run("Date: %s"%(datetime.datetime.now().strftime('%d %B %G'))).bold=True
document.tables[-1].cell(1,0).paragraphs[0].paragraph_format.space_after = Pt(1)
document.tables[-1].cell(2,0).text='Lorem ipsum'
document.tables[-1].cell(3,0).text='Signature of RA:\n'
document.tables[-1].cell(4,0).text="Supervisor's remarks\n"
document.tables[-1].cell(5,0).text="Date: %s"%(datetime.datetime.now().strftime('%d/%m/%Y'))
document.tables[-1].cell(6,0).text="Signature of supervisor:\n"

document.save("August Logbook 130.docx")

print('New entry added successfully.\n')
input('Press enter to close this window.')




